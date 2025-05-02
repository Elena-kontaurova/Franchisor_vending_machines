import tkinter as tk
from tkinter import PhotoImage, messagebox
from connect import Svodka, News, Torfavt, Kompany, AutorizRegus, \
    Otchet_torgov_avtomat
import random
from grafic import graf, svodk, graf_1, graf_2
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import openpyxl
import csv
from docx import Document


root = tk.Tk()
root.geometry('1124x700')
root.resizable(False, False)
root.title('Франчазер')
root.configure(background='#6b6b5e')


def get_otchet_torg_avt():
    c = Otchet_torgov_avtomat.select()
    return c


def get_svodka():
    svodka = Svodka.select()
    return svodka


def get_news():
    news = News.select()
    return news


def open_main_str(event=None):
    global kv1, hh
    ''' Личный кабинет главная'''
    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Главная',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=1050, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)
    kk = tk.Label(hh, text='Личный кабинет. Главная',
                  background='#c4cacf', font=('', 15))
    kk.place(x=20, y=20)
    kv1 = tk.Label(hh, width=32, height=12, background='white')
    kv1.place(x=20, y=60)
    pl_k1 = tk.Label(hh, width=32, height=3, background='#d1d8de')
    pl_k1.place(x=20, y=60)

    tex_avt = tk.Label(kv1,
                       background='white', width=35, height=9)
    tex_avt.place(x=0, y=50)

    kartinka = graf()
    uo = tk.Label(tex_avt, image=kartinka, background='white')
    uo.place(x=20, y=10)

    t1 = tk.Label(pl_k1, text='Эффективность сети', background='#d1d8de',
                  font=('', 11))
    t1.place(x=4, y=0)

    kv2 = tk.Label(hh, width=32, height=12, background='white')
    kv2.place(x=300, y=60)
    pl_k2 = tk.Label(hh, width=32, height=3, background='#d1d8de')
    pl_k2.place(x=300, y=60)

    s = svodk()

    t2 = tk.Label(pl_k2, text='Состояние сети', background='#d1d8de',
                  font=('', 11))
    t2.place(x=4, y=9)
    tex_avt = tk.Label(kv2, image=s,
                       background='white')
    tex_avt.place(x=40, y=50)

    kv3 = tk.Label(hh, width=32, height=12, background='white')
    kv3.place(x=580, y=60)
    pl_k3 = tk.Label(hh, width=32, height=3, background='#d1d8de')
    pl_k3.place(x=580, y=60)
    t3 = tk.Label(pl_k3, text='Сводка', background='#d1d8de',
                  font=('', 11))
    t3.place(x=4, y=9)
    svo = get_svodka()

    y_position = 52
    for i in svo:
        a = tk.Label(kv3, text=f'{i.name}', background='white',
                     font=('', 6))
        a.place(x=10, y=y_position)
        b = tk.Label(kv3, text=f'{i.price}', background='white',
                     font=('', 6))
        b.place(x=170, y=y_position)
        y_position += 15

    kv4 = tk.Label(hh, width=72, height=17, background='white')
    kv4.place(x=20, y=280)
    pl_k4 = tk.Label(hh, width=72, height=6, background='#d1d8de')
    pl_k4.place(x=20, y=280)
    t4 = tk.Label(pl_k4, text='Динамика продаж за последние 10 дней',
                  background='#d1d8de',
                  font=('', 11))

    kkk = tk.Label(kv4, background='white', width=70, height=10)
    kkk.place(x=0, y=100)

    t4.place(x=4, y=9)
    tete = tk.Label(pl_k4,
                    text='Данные по продажам с 1.03.2025 по 10.03.2025',
                    background='#d1d8de',
                    font=('', 7),
                    fg='#686b6e')

    def cd_blu(_):
        knop.config(background='#1870b8')
        knop1.config(background='white')
        hhk.config(background='#1870b8', fg='white')
        hhn.config(background='white', fg='black')

    def cd_hh(_):
        knop1.config(background='#1870b8')
        knop.config(background='white')
        hhn.config(background='#1870b8', fg='white')
        hhk.config(background='white', fg='black')

    pp = graf_1()
    lkj = tk.Label(kkk, image=pp, background='white')
    lkj.place(x=50, y=0)

    def lox():
        a = graf_1()
        j = tk.Label(kkk, image=a, background='white')
        j.place(x=50, y=0)

    def lox_2():
        a = graf_2()
        j = tk.Label(kkk, image=a, background='white')
        j.place(x=50, y=0)

    tete.place(x=4, y=37)
    knop = tk.Label(pl_k4, width=18, height=1, background='#d1d8de')
    knop.place(x=4, y=60)
    hhk = tk.Button(knop, text='По суммe', background='#1870b8', fg='white',
                    command=lox, width=18)
    hhk.place(x=0, y=0)
    knop.bind('<Button-1>', cd_blu)

    knop1 = tk.Label(pl_k4, width=18, height=1, background='#d1d8de')
    knop1.place(x=140, y=60)
    hhn = tk.Button(knop1, text='По количеству', background='white',
                    fg='black', width=18, command=lox_2)
    hhn.place(x=0, y=0)
    knop1.bind('<Button-1>', cd_hh)

    kv5 = tk.Label(hh, width=32, height=17, background='white')
    kv5.place(x=580, y=280)
    pl_k5 = tk.Label(hh, width=32, height=3, background='#d1d8de')
    pl_k5.place(x=580, y=280)
    t5 = tk.Label(pl_k5, text='Новости', background='#d1d8de',
                  font=('', 11))
    t5.place(x=4, y=9)

    news = get_news()
    y_posi = 53
    for i in news:
        c = tk.Label(kv5, text=f'{i.date}', background='white',
                     font=('', 6))
        c.place(x=10, y=y_posi)
        d = tk.Label(kv5, text=f'{i.text}', background='white',
                     font=('', 6))
        d.place(x=60, y=y_posi)
        y_posi += 30


def open_monik_str(_):
    ''' СТРАНИЦА МОНИТОР ТА'''
    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Монитор ТА',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=1025, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)

    www = tk.Label(hh, background='white', width=106, height=35)
    www.place(x=40, y=40)

    hh = tk.Label(www, background='#dbdbdb', width=180, height=10)
    hh.place(x=0, y=0)
    tete = tk.Label(www, text='380649 "ООО Торговые Автоматы"',
                    background='white')
    tete.place(x=16, y=156)
    teta = tk.Label(www,
                    text='Итого автоматов 9. Денег в автомата: 22460р',
                    background='white')
    teta.place(x=16, y=500)

    tex = tk.Label(hh, text='Монитор торговых автоматов',
                   background='#dbdbdb', fg='#0394fc',
                   font=('', 16))
    tex.place(x=10, y=0)
    tex_ob = tk.Label(hh, text='Общее состояние', background='#dbdbdb',)
    tex_ob.place(x=10, y=42)

    def button_pressed_o1():
        o1.config(relief='sunken')
        o1.config(background='#0c4519')

    o1 = tk.Button(hh, background='#06731d', text='',
                   relief='solid', border=1, height=2, width=5,
                   command=button_pressed_o1)
    o1.place(x=14, y=68)

    def button_pressed_o2():
        o2.config(relief='sunken')
        o2.config(background='#8f1b29')

    o2 = tk.Button(hh, background='#f20723', text='',
                   relief='solid', border=1, height=2, width=5,
                   command=button_pressed_o2)
    o2.place(x=54, y=68)

    def button_pressed_o3():
        o3.config(relief='sunken')
        o3.config(background='#0f517f')

    o3 = tk.Button(hh, background='#0394fc', text='',
                   relief='solid', border=1, height=2, width=5,
                   command=button_pressed_o3)
    o3.place(x=94, y=68)

    tex_pod = tk.Label(hh, text='Подключение', background='#dbdbdb')
    tex_pod.place(x=150, y=42)

    def button_pressed_p1():
        p1.config(relief='sunken')
        p1.config(background='#dce0e3')

    p1 = tk.Button(hh, background='white', text='Эк',
                   relief='solid', border=1, height=2, width=5,
                   fg='#0394fc', command=button_pressed_p1)
    p1.place(x=150, y=68)

    def button_pressed_p2():
        p2.config(relief='sunken')
        p2.config(background='#dce0e3')

    p2 = tk.Button(hh, background='white', text='MDV',
                   relief='solid', border=1, height=2, width=5,
                   fg='#0394fc', command=button_pressed_p2)
    p2.place(x=190, y=68)

    def button_pressed_p3():
        p3.config(relief='sunken')
        p3.config(background='#dce0e3')

    p3 = tk.Button(hh, background='white', text='EXE',
                   relief='solid', border=1, height=2, width=5,
                   fg='#0394fc', command=button_pressed_p3)
    p3.place(x=230, y=68)

    tex_dop = tk.Label(hh, text='Дополнительные статусы',
                       background='#dbdbdb')
    tex_dop.place(x=290, y=42)

    def button_pressed_d1():
        d1.config(relief='sunken')
        d1.config(background='#dce0e3')

    d1 = tk.Button(hh, background='white', text='1',
                   relief='solid', border=1, height=2, width=5,
                   fg='#f20723', command=button_pressed_d1)
    d1.place(x=290, y=68)

    def button_pressed_d2():
        d2.config(relief='sunken')
        d2.config(background='#dce0e3')

    d2 = tk.Button(hh, background='white', text='2',
                   relief='solid', border=1, height=2, width=5,
                   fg='#f20723', command=button_pressed_d2)
    d2.place(x=330, y=68)

    def button_pressed_d3():
        d3.config(relief='sunken')
        d3.config(background='#dce0e3')

    d3 = tk.Button(hh, background='white', text='3',
                   relief='solid', border=1, height=2, width=5,
                   fg='#0394fc', command=button_pressed_d3)
    d3.place(x=370, y=68)

    def button_pressed_d4():
        d4.config(relief='sunken')
        d4.config(background='#dce0e3')

    d4 = tk.Button(hh, background='white', text='4',
                   relief='solid', border=1, height=2, width=5,
                   fg='#0394fc', command=button_pressed_d4)
    d4.place(x=410, y=68)

    def button_pressed_d5():
        d5.config(relief='sunken')
        d5.config(background='#dce0e3')

    d5 = tk.Button(hh, background='white', text='5',
                   relief='solid', border=1, height=2, width=5,
                   fg='#0394fc', command=button_pressed_d5)
    d5.place(x=450, y=68)

    tex_sor = tk.Label(hh, text='Сортировка', background='#dbdbdb')
    tex_sor.place(x=510, y=42)
    a1 = tk.Label(hh, background='white', text='По состоянию ТА',
                  relief='solid', border=1, height=2, width=25,
                  fg='#999696')
    a1.place(x=510, y=68)

    def clearing_the_sorting():
        o1.config(relief='solid')
        o1.config(background='#06731d')

        o2.config(relief='solid')
        o2.config(background='#f20723')

        o3.config(relief='solid')
        o3.config(background='#0394fc')

        p1.config(relief='solid')
        p1.config(background='white')

        p2.config(relief='solid')
        p2.config(background='white')

        p3.config(relief='solid')
        p3.config(background='white')

        d1.config(relief='solid')
        d1.config(background='white')

        d2.config(relief='solid')
        d2.config(background='white')

        d3.config(relief='solid')
        d3.config(background='white')

        d4.config(relief='solid')
        d4.config(background='white')

        d5.config(relief='solid')
        d5.config(background='white')

    pri = tk.Button(hh, text='Применить', background='#0394fc',
                    fg='white', width=19)
    pri.place(x=13, y=115)
    ohi = tk.Button(hh, text='Очистить', background='white',
                    fg='black', width=19, command=clearing_the_sorting)
    ohi.place(x=170, y=115)

    table = tk.Label(www, background='white', border=1,
                     relief='sunken', width=100, height=21)
    table.place(x=20, y=180)
    oglav = tk.Label(table, background='#dbdbdb', border=1,
                     relief='sunken', width=100, height=3)
    oglav.place(x=0, y=0)

    num = tk.Label(oglav, text='#', background='#dbdbdb')
    num.place(x=15, y=15)
    tor_avt = tk.Label(oglav, text='Торговый автомат', background='#dbdbdb')
    tor_avt.place(x=48, y=15)
    svaz = tk.Label(oglav, text='Связь', background='#dbdbdb')
    svaz.place(x=174, y=15)
    zagr = tk.Label(oglav, text='Загрузка', background='#dbdbdb')
    zagr.place(x=234, y=15)
    dene = tk.Label(oglav, text='Денеж. средства',
                    background='#dbdbdb')
    dene.place(x=300, y=15)
    sob = tk.Label(oglav, text='События', background='#dbdbdb')
    sob.place(x=420, y=15)
    oboru = tk.Label(oglav, text='Оборудование', background='#dbdbdb')
    oboru.place(x=494, y=15)
    info = tk.Label(oglav, text='Информация', background='#dbdbdb')
    info.place(x=584, y=15)
    dop = tk.Label(oglav, text='Доп.', background='#dbdbdb')
    dop.place(x=670, y=15)

    p = tk.Label(table, background='#f7f7fa', border=1,
                 relief='sunken', width=6, height=20)
    p.place(x=0, y=48)

    vt = tk.Label(table, background='#f7f7fa', border=1,
                  relief='sunken', width=16, height=18)
    vt.place(x=46, y=48)

    tr = tk.Label(table, background='#f7f7fa', border=1,
                  relief='sunken', width=14, height=18)
    tr.place(x=160, y=48)

    che = tk.Label(table, background='#f7f7fa', border=1,
                   relief='sunken', width=10, height=18)
    che.place(x=230, y=48)

    pat = tk.Label(table, background='#f7f7fa', border=1,
                   relief='sunken', width=15, height=18)
    pat.place(x=300, y=48)

    she = tk.Label(table, background='#f7f7fa', border=1,
                   relief='sunken', width=13, height=18)
    she.place(x=405, y=48)

    sem = tk.Label(table, background='#f7f7fa', border=1,
                   relief='sunken', width=13, height=18)
    sem.place(x=495, y=48)

    vos = tk.Label(table, background='#f7f7fa', border=1,
                   relief='sunken', width=13, height=18)
    vos.place(x=585, y=48)

    dev = tk.Label(table, background='#f7f7fa', border=1,
                   relief='sunken', width=5, height=18)
    dev.place(x=665, y=48)

    pp = 80
    oo = 60
    gg = 90
    hh = 82
    ggg = 98
    l1 = 65
    l2 = 93
    za = 76
    kk = 70
    for i in range(4):
        tk.Label(table, text=f'{i + 1}',
                 background='#f7f7fa').place(x=20, y=pp)
        tk.Label(table, text='903823 - "БЦ Московский"',
                 background='#f7f7fa', justify='center',
                 wraplength=75).place(x=80, y=oo)
        tk.Label(table, text='Unicon Rosso'
                 '\n Суворово 21',
                 background='#f7f7fa', justify='center',
                 wraplength=75, fg='#999696',
                 font=('', 8)).place(x=80, y=gg)
        tk.Label(table, border=1, relief='solid',
                 background='#0394fc', width=3,
                 height=1).place(x=53, y=pp)
        tk.Label(table, text='T2',
                 background='#f7f7fa', font=('', 12)).place(x=182, y=oo)
        tk.Label(table, text='21 р',
                 background='#f7f7fa',
                 font=('', 9)).place(x=180, y=hh)
        tk.Label(table, text='23:45',
                 background='#f7f7fa',
                 font=('', 9, 'bold')).place(x=175, y=ggg)
        tk.Label(table, text='обща.',
                 background='#07eafa').place(x=245, y=l1)
        tk.Label(table, text='мини.',
                 background='#07eafa').place(x=245, y=gg)

        tk.Label(table, text='0 р.',
                 background='#f7f7fa',
                 font=('', 9, 'bold')).place(x=320, y=oo)
        tk.Label(table, text='7820 р.',
                 background='#f7f7fa',
                 font=('', 9, 'bold')).place(x=320, y=za)
        tk.Label(table, text='11 345 р.',
                 background='#f7f7fa',
                 font=('', 9, 'bold')).place(x=320, y=l2)

        tk.Label(table, text='11 мин. назад',
                 background='#f7f7fa').place(x=410, y=oo)
        tk.Label(table, text='2 дня назад',
                 background='#f7f7fa').place(x=410, y=za)
        tk.Label(table, text='2 дня назад',
                 background='#f7f7fa').place(x=410, y=l2)

        tk.Label(table, text='247',
                 background='#f7f7fa').place(x=673, y=pp)

        tk.Label(table, border=1, relief='solid',
                 background='#06731d',
                 width=3, height=1).place(x=513, y=kk)
        tk.Label(table, border=1, relief='solid',
                 background='#06731d',
                 width=3, height=1).place(x=543, y=kk)
        tk.Label(table, border=1, relief='solid',
                 background='#06731d',
                 width=3, height=1).place(x=528, y=l2)

        tk.Label(table, border=1, relief='solid',
                 background='white', fg='#0394fc', text='EXE',
                 width=3, height=1).place(x=600, y=kk)
        tk.Label(table, border=1, relief='solid',
                 background='white', fg='#0394fc', text='MV',
                 width=3, height=1).place(x=630, y=kk)
        tk.Label(table, border=1, relief='solid',
                 background='white', fg='#0394fc', text='ZK',
                 width=3, height=1).place(x=615, y=l2)

        pp += 60
        oo += 62
        gg += 62
        hh += 62
        ggg += 62
        l1 += 62
        l2 += 62
        za += 62
        kk += 62


def open_det_otc_str_1(_):
    ''' детальный отчет первый'''
    ff = tk.Label(root, background='#060a0d', width=70, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Детальные отчеты/ Отчет 1',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=905, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)

    lable_pod_formy = tk.Label(hh, background='#e1e5e8', width=115,
                               height=37)
    lable_pod_formy.place(x=12, y=20)

    tabel = tk.Label(lable_pod_formy, background='#fcfcfc',
                     width=60, height=15)
    tabel.place(x=40, y=60)

    tk.Label(lable_pod_formy, text='Отчет о торговые автоматах',
             background='#e1e5e8', font=('', 15)).place(x=50, y=20)

    otchet = get_otchet_torg_avt()

    def export_to_pdf(otchet):
        ''' метод для экспорта в pdf файл'''
        pdf_file = "report_torgov_avt.pdf"
        c = canvas.Canvas(pdf_file, pagesize=letter)
        width, height = letter

        pdfmetrics.registerFont(TTFont('Arial', 'Arial.ttf'))
        c.setFont('Arial', 12)

        y_position = height - 40

        for i in otchet:
            c.drawString(50, y_position,
                         f'Всего автоматов: {i.itigo_avtomatov}')
            y_position -= 20
            c.drawString(50, y_position,
                         f'Использованно автоматов: {i.uspolzuen}')
            y_position -= 20
            c.drawString(50, y_position, f'Свободно автоматов: {i.svobodno}')
            y_position -= 20
            c.drawString(50, y_position, f'Работает автоматов: {i.rabotaet}')
            y_position -= 20
            c.drawString(50, y_position,
                         f'Не работает автоматов: {i.ne_rabotaey}')
            y_position -= 20
            c.drawString(50, y_position,
                         f'Требуется обслуживание: {i.trebue_obsluch}')
            y_position -= 20
            c.drawString(50, y_position, f'Требуется проверка: {i.provetka}')
            y_position -= 20

            if y_position < 40:
                c.showPage()
                y_position = height - 40

        c.save()
        print(f"PDF файл '{pdf_file}' успешно сохранен.")

    def export_to_excel(otchet):
        workbook = openpyxl.Workbook()
        sheet = workbook.active

        headers = ["Всего автоматов", "Использованно автоматов",
                   "Свободно автоматов", "Работает автоматов",
                   "Не работает автоматов", "Требуется обслуживание",
                   "Требуется проверка"]
        sheet.append(headers)

        for i in otchet:
            row = [
                i.itigo_avtomatov,
                i.uspolzuen,
                i.svobodno,
                i.rabotaet,
                i.ne_rabotaey,
                i.trebue_obsluch,
                i.provetka
            ]
            sheet.append(row)

        excel_file = "report_torg_avt.xlsx"
        workbook.save(excel_file)
        print(f"Excel файл '{excel_file}' успешно сохранен.")

    def export_to_csv(otchet):
        csv_file = "report_torg_avt.csv"
        with open(csv_file, mode='w', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)

            headers = ["Всего автоматов", "Использованно автоматов",
                       "Свободно автоматов", "Работает автоматов",
                       "Не работает автоматов", "Требуется обслуживание",
                       "Требуется проверка"]
            writer.writerow(headers)

            for i in otchet:
                row = [
                    i.itigo_avtomatov,
                    i.uspolzuen,
                    i.svobodno,
                    i.rabotaet,
                    i.ne_rabotaey,
                    i.trebue_obsluch,
                    i.provetka
                ]
                writer.writerow(row)

        print(f"CSV файл '{csv_file}' успешно сохранен.")

    def export_to_word(otchet):
        doc = Document()

        doc.add_heading('Отчет торговые автоматы', level=1)

        for i in otchet:
            doc.add_paragraph(f'Всего автоматов: {i.itigo_avtomatov}')
            doc.add_paragraph(f'Использованно автоматов: {i.uspolzuen}')
            doc.add_paragraph(f'Свободно автоматов: {i.svobodno}')
            doc.add_paragraph(f'Работает автоматов: {i.rabotaet}')
            doc.add_paragraph(f'Не работает автоматов: {i.ne_rabotaey}')
            doc.add_paragraph(f'Требуется обслуживание: {i.trebue_obsluch}')
            doc.add_paragraph(f'Требуется проверка: {i.provetka}')
            doc.add_paragraph('')

        word_file = "report_torg_avt.docx"
        doc.save(word_file)
        print(f"Word файл '{word_file}' успешно сохранен.")

    for i in otchet:
        tk.Label(lable_pod_formy, text='Всего автоматов: ',
                 background='#fcfcfc', font=('', 13)).place(x=50, y=80)
        tk.Label(lable_pod_formy, text=f'{i.itigo_avtomatov}',
                 background='#fcfcfc', font=('', 13)).place(x=290, y=80)

        tk.Label(lable_pod_formy, text='Использованно автоматов: ',
                 background='#fcfcfc', font=('', 13)).place(x=50, y=110)
        tk.Label(lable_pod_formy, text=f'{i.uspolzuen}',
                 background='#fcfcfc', font=('', 13)).place(x=290, y=110)

        tk.Label(lable_pod_formy, text='Свободно автоматов: ',
                 background='#fcfcfc', font=('', 13)).place(x=50, y=140)
        tk.Label(lable_pod_formy, text=f'{i.svobodno}',
                 background='#fcfcfc', font=('', 13)).place(x=290, y=140)

        tk.Label(lable_pod_formy, text='Работает автоматов: ',
                 background='#fcfcfc', font=('', 13)).place(x=50, y=170)
        tk.Label(lable_pod_formy, text=f'{i.rabotaet}',
                 background='#fcfcfc', font=('', 13)).place(x=290, y=170)

        tk.Label(lable_pod_formy, text='Не работает автоматов: ',
                 background='#fcfcfc', font=('', 13)).place(x=50, y=200)
        tk.Label(lable_pod_formy, text=f'{i.ne_rabotaey}',
                 background='#fcfcfc', font=('', 13)).place(x=290, y=200)

        tk.Label(lable_pod_formy, text='Требуется обслуживание: ',
                 background='#fcfcfc', font=('', 13)).place(x=50, y=230)
        tk.Label(lable_pod_formy, text=f'{i.trebue_obsluch}',
                 background='#fcfcfc', font=('', 13)).place(x=290, y=230)

        tk.Label(lable_pod_formy, text='Требуется проверка: ',
                 background='#fcfcfc', font=('', 13)).place(x=50, y=260)
        tk.Label(lable_pod_formy, text=f'{i.provetka}',
                 background='#fcfcfc', font=('', 13)).place(x=290, y=260)

        button_pdf = tk.Button(lable_pod_formy, text='Экспортировать в pdf',
                               width=20, background='#0e8ae3', fg='white',
                               command=lambda: export_to_pdf(otchet))
        button_pdf.place(x=50, y=320)

        button_excel = tk.Button(lable_pod_formy,
                                 text='Экспортировать в excel',
                                 width=20, background='#0e8ae3', fg='white',
                                 command=lambda: export_to_excel(otchet))
        button_excel.place(x=50, y=360)

        button_csv = tk.Button(lable_pod_formy, text='Экспортировать в csv',
                               background='#0e8ae3', width=20, fg='white',
                               command=lambda: export_to_csv(otchet))
        button_csv.place(x=50, y=400)

        button_word = tk.Button(lable_pod_formy, text='Экспортировать в .txt',
                                background='#0e8ae3', width=20, fg='white',
                                command=lambda: export_to_word(otchet))
        button_word.place(x=50, y=440)


def open_det_otc_str_2(_):
    ''' детальные отчеты - второй'''
    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Детальные отчеты/ Отчет 2',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=905, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)


def open_det_otc_str_3(_):
    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=900, y=50)
    gg = tk.Label(root, text='Детальные отчеты / Отчет 3',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=905, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)


def open_ychet_1(_):
    ff = tk.Label(root, background='#060a0d', width=70, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Учет ТМЦ / Учет 1',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=975, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)


def open_ychet_2(_):
    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Учет ТМЦ / Учет 2',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=975, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)


def open_ychet_3(_):
    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Учет ТМЦ/ Учет 3',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=975, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)


def get_tor():
    torg = Torfavt.select()
    return torg


def get_komp():
    kom = Kompany.select()
    return kom


def open_torg_avt(_):
    ''' Администрирование / торговые автоматы'''

    def create_zap():
        ff = tk.Toplevel()
        ff.geometry('300x500')
        ff.title('Создание торгового автомата')

        la = tk.Label(ff, background='#c4cacf', width=300, height=3)
        la.place(x=0, y=0)
        oglav = tk.Label(ff, text='Создание торгового аппарата',
                         font=('', 15), fg='#1489cc', background='#c4cacf')
        oglav.place(x=10, y=10)

        name = tk.Label(ff, text='Название автомата:')
        name.place(x=20, y=60)
        name1 = tk.Entry(ff)
        name1.place(x=140, y=60)

        model = tk.Label(ff, text='Модель:')
        model.place(x=20, y=90)
        m = tk.Entry(ff)
        m.place(x=140, y=90)

        kompany = tk.Label(ff, text='Компания:')
        kompany.place(x=20, y=120)
        k = tk.Entry(ff)
        k.place(x=140, y=120)

        modem = tk.Label(ff, text='Модем:')
        modem.place(x=20, y=150)
        m1 = tk.Entry(ff)
        m1.place(x=140, y=150)

        adress = tk.Label(ff, text='Адерес/Место:')
        adress.place(x=20, y=180)
        a = tk.Entry(ff)
        a.place(x=140, y=180)

        word = tk.Label(ff, text='В работе с:')
        word.place(x=20, y=210)
        w = tk.Entry(ff)
        w.place(x=140, y=210)

        deist = tk.Label(ff, text='Действия:')
        deist.place(x=20, y=240)
        d = tk.Entry(ff)
        d.place(x=140, y=240)

        def gr_zz():
            name = name1.get()
            model = m.get()
            kompany = k.get()
            modem = m1.get()
            adress = a.get()
            word = w.get()
            deist = d.get()

            _ = Torfavt.create(
                name=name,
                model=model,
                kompany=kompany,
                modem=modem,
                adress=adress,
                word=word,
                deist=deist
            )
            messagebox.showinfo('Успех', 'Запись успещно создана')
            ff.destroy()

        cre = tk.Button(ff, text='Создать', background='#1489cc',
                        width=13, fg='white', command=gr_zz)
        cre.place(x=20, y=280)

        ff.mainloop()

    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Администирование / Торговые автоматы',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=800, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)
    kvad = tk.Label(hh, width=100, height=34, background='white')
    kvad.place(x=70, y=30)
    hh = tk.Label(kvad, background='#ededed', width=100, height=3)
    hh.place(x=0, y=0)
    tt = tk.Label(hh, text='Торговые автоматы', background='#ededed',
                  fg='#1489cc', font=('', 12))
    tt.place(x=10, y=5)
    df = tk.Label(hh, text='Всего найдено 9 записей', background='#ededed',
                  fg='black', font=('', 7))
    df.place(x=12, y=25)
    kn_create = tk.Button(hh, text='Добавить', background='#ededed',
                          fg='black', command=create_zap)
    kn_create.place(x=610, y=15)

    table = tk.Label(kvad, width=97, height=29, background='#adb4b8')
    table.place(x=10, y=60)
    t_1 = tk.Label(table,
                   text='ID', background='#adb4b8')
    t_1.place(x=30, y=10)
    t_2 = tk.Label(table,
                   text='Навазание \n автомата',
                   background='#adb4b8')
    t_2.place(x=90, y=5)
    t_3 = tk.Label(table, text='Модель', background='#adb4b8')
    t_3.place(x=190, y=10)
    t_4 = tk.Label(table, text='Компания', background='#adb4b8')
    t_4.place(x=270, y=10)
    t_5 = tk.Label(table, text='Модем', background='#adb4b8')
    t_5.place(x=360, y=10)
    t_6 = tk.Label(table, text='Адресс/Место', background='#adb4b8')
    t_6.place(x=430, y=10)
    t_7 = tk.Label(table, text='В работе \n с', background='#adb4b8')
    t_7.place(x=530, y=5)
    t_8 = tk.Label(table, text='Действия', background='#adb4b8')
    t_8.place(x=600, y=10)

    def update_torg_avt(item, name, model, kompany, modem, adress, word):
        aga = Torfavt.get(Torfavt.id == item.id)
        aga.name = name
        aga.model = model
        aga.kompany = kompany
        aga.modem = modem
        aga.adress = adress
        aga.word = word
        aga.save()

    def redak_torg_avt(item):
        mm = tk.Toplevel()
        mm.geometry('500x300')
        mm.title('Редактирование торгового автомата')

        a = tk.Label(mm, text='Название автомата:')
        a.place(x=20, y=20)
        name = tk.Entry(mm)
        name.insert(0, item.name)
        name.place(x=20, y=45)

        b = tk.Label(mm, text='Модель:')
        b.place(x=180, y=20)

        model = tk.Entry(mm)
        model.insert(0, item.model)
        model.place(x=180, y=45)

        c = tk.Label(mm, text='Компания:')
        c.place(x=20, y=85)

        komp = tk.Entry(mm, width=40)
        komp.insert(0, item.kompany)
        komp.place(x=20, y=110)

        d = tk.Label(mm, text='Модем:')
        d.place(x=280, y=85)

        modem = tk.Entry(mm)
        modem.insert(0, item.modem)
        modem.place(x=280, y=110)

        e = tk.Label(mm, text='Адресс/Место:')
        e.place(x=20, y=145)

        adress = tk.Entry(mm, width=30)
        adress.insert(0, item.adress)
        adress.place(x=20, y=170)

        f = tk.Label(mm, text='В работе с')
        f.place(x=230, y=145)

        word = tk.Entry(mm)
        word.insert(0, item.word)
        word.place(x=230, y=170)

        kn = tk.Label(mm, text='Сохранить', width=25,
                      background='#1489CC', fg='white')

        kn = tk.Button(mm, text='Сохранить', width=25,
                       background='#1489CC', fg='white')
        kn.place(x=20, y=220)

        def update_info():
            new_name = name.get()
            new_model = model.get()
            new_komp = komp.get()
            new_modem = modem.get()
            new_adress = adress.get()
            new_word = word.get()
            update_torg_avt(item, new_name, new_model, new_komp,
                            new_modem, new_adress, new_word)

            messagebox.showinfo("Успех", "Данные успешно сохранены!")
            mm.destroy()

        kn.config(command=update_info)

        mm.mainloop()

    def delete_torg_avt(item):
        mm = tk.Toplevel()
        mm.geometry('250x100')
        mm.title('Удаление торгового автомата')

        def del_zap():
            Torfavt.delete().where(Torfavt.id == item.id).execute()
            messagebox.showinfo('Успех', 'Данные успешно удаленны')
            mm.destroy()

        def del_fram():
            mm.destroy()

        a = tk.Label(mm, text='Вы действительно хотите удалить запись?')
        a.place(x=10, y=10)
        yes = tk.Button(mm, text='Да', background='#f00c36', fg='white',
                        width=13, command=del_zap)
        yes.place(x=10, y=50)
        no = tk.Button(mm, text='Нет', background='#0c14f0', fg='white',
                       width=13, command=del_fram)
        no.place(x=140, y=50)
        mm.mainloop()

    torg = get_tor()
    lox = 50
    for i in torg:
        asas = '#d9dadb'
        if i.id % 2 == 0:
            asas = '#ededed'
        dfd = tk.Label(table, width=97, height=2, background=asas)
        dfd.place(x=0, y=lox)
        lox += 40
        h_1 = tk.Label(dfd, text=f'{i.id}', background=asas)
        h_1.place(x=30, y=5)
        h_2 = tk.Label(dfd, text=f'{i.name}', background=asas,
                       wraplength=90, justify='center', fg='#1489CC')
        h_2.place(x=100, y=5)
        h_3 = tk.Label(dfd, text=f'{i.model}', background=asas,
                       wraplength=70, justify='center', font=('', 8))
        h_3.place(x=190, y=0)
        h_4 = tk.Label(dfd, text=f'{i.kompany}', background=asas,
                       wraplength=80, justify='center', font=('', 6),
                       fg='#1489CC')
        h_4.place(x=260, y=0)
        h_5 = tk.Label(dfd, text=f'{i.modem}', background=asas)
        h_5.place(x=350, y=5)
        h_6 = tk.Label(dfd, text=f'{i.adress}', background=asas,
                       wraplength=70, justify='center', font=('', 8))
        h_6.place(x=440, y=0)
        h_7 = tk.Label(dfd, text=f'{i.word}', background=asas)
        h_7.place(x=530, y=5)
        h_8 = tk.Label(dfd, text='R', background=asas,
                       fg='#1489CC')
        h_8.place(x=610, y=5)
        h_8.bind('<Button-1>', lambda event, item=i: redak_torg_avt(item))

        h_9 = tk.Label(dfd, text='D', background=asas,
                       fg='#1489CC')
        h_9.place(x=630, y=5)
        h_9.bind('<Button-1>', lambda event, item=i: delete_torg_avt(item))

        h_0 = tk.Label(dfd, text='C', background=asas,
                       fg='#1489CC')
        h_0.place(x=650, y=5)


def open_komp(_):
    ''' Администрирование / Компании '''

    def create_komp():
        mm = tk.Toplevel()
        mm.geometry('500x200')
        mm.title('Создание компании')

        a = tk.Label(mm, text='Название:')
        a.place(x=30, y=10)
        nana = tk.Entry(mm)
        nana.place(x=30, y=30)

        b = tk.Label(mm, text='Вышестоящая:')
        b.place(x=190, y=10)
        veve = tk.Entry(mm)
        veve.insert(0, 'ООО Торговые автоматы')
        veve.place(x=190, y=30)

        c = tk.Label(mm, text='Адрес:')
        c.place(x=350, y=10)
        adad = tk.Entry(mm)
        adad.place(x=350, y=30)

        d = tk.Label(mm, text='Контакты:')
        d.place(x=30, y=60)
        koko = tk.Entry(mm)
        koko.place(x=30, y=80)

        e = tk.Label(mm, text='В работе с:')
        e.place(x=190, y=60)
        wowo = tk.Entry(mm)
        wowo.place(x=190, y=80)

        f = tk.Label(mm, text='Примечание:')
        f.place(x=350, y=60)
        prpr = tk.Entry(mm)
        prpr.place(x=350, y=80)

        def yes_button():
            name = nana.get()
            veshe = veve.get()
            adres = adad.get()
            kontak = koko.get()
            work = wowo.get()
            prim = prpr.get()

            _ = Kompany.create(
                name=name,
                veshe=veshe,
                adres=adres,
                kontak=kontak,
                work=work,
                prim=prim
            )
            messagebox.showinfo('Успех', 'Компания успешно создана')
            mm.destroy()

        def no_button():
            mm.destroy()

        yes = tk.Button(mm, text='Добавить', width=13,
                        background='#0c14f0', fg='white',
                        command=yes_button)
        yes.place(x=30, y=120)
        no = tk.Button(mm, text='Отменить', background='#c2c2c4',
                       fg='black', width=13, command=no_button)
        no.place(x=150, y=120)

    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Администирование / Компании',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=870, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)
    kvad = tk.Label(hh, width=100, height=34, background='white')
    kvad.place(x=70, y=30)
    hh = tk.Label(kvad, background='#ededed', width=100, height=3)
    hh.place(x=0, y=0)
    tt = tk.Label(hh, text='Cписок компаний', background='#ededed',
                  fg='#1489cc', font=('', 12))
    tt.place(x=10, y=5)
    df = tk.Label(hh, text='Всего найдено 9 записей', background='#ededed',
                  fg='black', font=('', 7))
    df.place(x=12, y=25)
    kn_create = tk.Button(hh, text='Добавить', background='#ededed',
                          fg='black', command=create_komp)
    kn_create.place(x=610, y=15)

    table = tk.Label(kvad, width=97, height=29, background='#adb4b8')
    table.place(x=10, y=60)
    t_2 = tk.Label(table,
                   text='Навазание',
                   background='#adb4b8')
    t_2.place(x=40, y=10)
    t_3 = tk.Label(table, text='Вышестоящая', background='#adb4b8')
    t_3.place(x=150, y=10)
    t_4 = tk.Label(table, text='Адрес', background='#adb4b8')
    t_4.place(x=290, y=10)
    t_5 = tk.Label(table, text='Контакты', background='#adb4b8')
    t_5.place(x=370, y=10)
    t_6 = tk.Label(table, text='В работе с', background='#adb4b8')
    t_6.place(x=460, y=10)
    t_7 = tk.Label(table, text='Действия', background='#adb4b8')
    t_7.place(x=560, y=10)

    def update_komp(item):
        mm = tk.Toplevel()
        mm.geometry('500x200')
        mm.title('Редактирование компании')

        a = tk.Label(mm, text='Название:')
        a.place(x=10, y=10)
        f1 = tk.Entry(mm)
        f1.insert(0, item.name)
        f1.place(x=10, y=35)

        b = tk.Label(mm, text='Вышестоящие:')
        b.place(x=160, y=10)
        f6 = tk.Entry(mm)
        f6.insert(0, item.veshe)
        f6.place(x=160, y=35)

        c = tk.Label(mm, text='Адресс:')
        c.place(x=310, y=10)
        f2 = tk.Entry(mm)
        f2.insert(0, item.adres)
        f2.place(x=310, y=35)

        d = tk.Label(mm, text='Контакты:')
        d.place(x=10, y=70)
        f3 = tk.Entry(mm)
        f3.insert(0, item.kontak)
        f3.place(x=10, y=95)

        e = tk.Label(mm, text='В работе с:')
        e.place(x=160, y=70)
        f4 = tk.Entry(mm)
        f4.insert(0, item.work)
        f4.place(x=160, y=95)

        def update_kk(item, name, veshe, adres, kontak,
                      work):
            kom = Kompany.get(Kompany.id == item.id)
            kom.name = name
            kom.veshe = veshe
            kom.adres = adres
            kom.kontak = kontak
            kom.work = work
            kom.save()

        def suda():
            name = f1.get()
            veshe = f6.get()
            adres = f2.get()
            kontak = f3.get()
            work = f4.get()

            update_kk(item, name, veshe, adres, kontak, work)
            messagebox.showinfo('Успех', 'Изменения сохранены')
            mm.destroy()

        kn = tk.Button(mm, text='Сохранить', width=20,
                       background='#1489CC', fg='white',
                       command=suda)
        kn.place(x=10, y=140)

    def del_komp(item):
        mm = tk.Toplevel()
        mm.geometry('300x100')
        mm.title('Удаление компании')

        def no_but():
            mm.destroy()

        def yes_but():
            Kompany.delete().where(Kompany.id == item.id).execute()
            messagebox.showinfo('Успех!', 'Данные успешно удаленны')
            mm.destroy()

        vop = tk.Label(mm, text='Вы действительно хотите удалить "Компанию"?')
        vop.place(x=10, y=10)
        yes = tk.Button(mm, text='Да', width=13,
                        background='#0c14f0', fg='white',
                        command=yes_but)
        yes.place(x=10, y=50)
        no = tk.Button(mm, text='Нет', background='#c2c2c4',
                       fg='black', width=13, command=no_but)
        no.place(x=130, y=50)

    kom = get_komp()
    lox = 50
    for i in kom:
        asas = '#d9dadb'
        if i.id % 2 == 0:
            asas = '#ededed'
        dfd = tk.Label(table, width=97, height=2, background=asas)
        dfd.place(x=0, y=lox)
        lox += 40

        h_2 = tk.Label(dfd, text=f'{i.name}', background=asas,
                       wraplength=90, justify='center', fg='#1489CC')
        h_2.place(x=40, y=0)

        h_3 = tk.Label(dfd, text=f'{i.veshe}', background=asas,
                       wraplength=70, justify='center', font=('', 7))
        h_3.place(x=160, y=0)

        h_4 = tk.Label(dfd, text=f'{i.adres}', background=asas,
                       wraplength=80, justify='center', font=('', 9),
                       fg='#1489CC')
        h_4.place(x=280, y=0)

        h_5 = tk.Label(dfd, text=f'{i.kontak}', background=asas)
        h_5.place(x=370, y=5)

        h_6 = tk.Label(dfd, text=f'{i.work}', background=asas,
                       wraplength=70, justify='center', font=('', 10))
        h_6.place(x=460, y=5)

        h_8 = tk.Label(dfd, text='R', background=asas,
                       fg='#1489CC')
        h_8.place(x=580, y=5)
        h_8.bind('<Button-1>', lambda event, item=i: update_komp(item))

        h_9 = tk.Label(dfd, text='D', background=asas,
                       fg='#1489CC')
        h_9.place(x=600, y=5)
        h_9.bind('<Button-1>', lambda event, item=i: del_komp(item))

        h_0 = tk.Label(dfd, text='C', background=asas,
                       fg='#1489CC')
        h_0.place(x=620, y=5)


def open_polsov(_):
    ''' Администирование / Пользователи'''
    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Администирование / Пользователи',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=850, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)


def open_modem(_):
    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Администирование / Модемы',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=890, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)


def open_dop(_):
    ff = tk.Label(root, background='#060a0d', width=100, height=10)
    ff.place(x=800, y=50)
    gg = tk.Label(root, text='Администирование / Дополнительные',
                  background='#060a0d', fg='#c4cacf',
                  font=('', 12))
    gg.place(x=830, y=60)
    hh = tk.Label(root, width=200, height=200,
                  background='#c4cacf')
    hh.place(x=285, y=100)


def open_glavna():
    def open_pod_menu_detot(_):
        global atsosite
        lkj.place(x=10, y=380)
        sdf.place(x=10, y=440)
        hhhhg.place(x=65, y=390)
        hhhhh.place(x=65, y=450)
        ff2.place(x=250, y=390)
        ff3.place(x=250, y=450)
        ff1.config(image=soss)
        atsosite = tk.Label(root, background='#0b1217',
                            width=40, height=5)
        atsosite.place(x=0, y=300)
        tt = tk.Label(atsosite, text='Отчет торговые автоматы',
                      background='#0b1217',
                      fg='white', font=('', 10))
        tt.place(x=30, y=10)
        tt.bind('<Button-1>', open_det_otc_str_1)
        tg = tk.Label(atsosite, text='Отчет мониторы', background='#0b1217',
                      fg='white', font=('', 10))
        tg.place(x=30, y=30)
        tg.bind('<Button-1>', open_det_otc_str_2)
        tf = tk.Label(atsosite, text='Отчет 3', background='#0b1217',
                      fg='white', font=('', 10))
        tf.place(x=30, y=50)
        tf.bind('<Button-1>', open_det_otc_str_3)

    def open_pod_menu_ychet(_):
        global ststs
        sdf.place(x=10, y=440)
        hhhhh.place(x=65, y=450)
        ff3.place(x=250, y=450)
        ff2.config(image=soss)
        ststs = tk.Label(root, background='#0b1217',
                         width=40, height=5)
        ststs.place(x=0, y=350)
        tt = tk.Label(ststs, text='Учет 1', background='#0b1217',
                      fg='white', font=('', 10))
        tt.place(x=30, y=10)
        tt.bind('<Button-1>', open_ychet_1)
        tg = tk.Label(ststs, text='Учет 2', background='#0b1217',
                      fg='white', font=('', 10))
        tg.place(x=30, y=30)
        tg.bind('<Button-1>', open_ychet_2)
        tf = tk.Label(ststs, text='Учет 3', background='#0b1217',
                      fg='white', font=('', 10))
        tf.place(x=30, y=50)
        tf.bind('<Button-1>', open_ychet_3)

    def open_pod_menu_admin(_):
        global ssss
        ff3.config(image=soss)
        ssss = tk.Label(root, background='#0b1217',
                        width=40, height=10)
        ssss.place(x=0, y=410)
        tt = tk.Label(ssss, text='Торговые автоматы', background='#0b1217',
                      fg='white', font=('', 10))
        tt.place(x=30, y=10)
        tt.bind('<Button-1>', open_torg_avt)
        tg = tk.Label(ssss, text='Компании', background='#0b1217',
                      fg='white', font=('', 10))
        tg.place(x=30, y=30)
        tg.bind('<Button-1>', open_komp)
        tf = tk.Label(ssss, text='Пользователи', background='#0b1217',
                      fg='white', font=('', 10))
        tf.place(x=30, y=50)
        tf.bind('<Button-1>', open_polsov)
        hh = tk.Label(ssss, text='Модемы', background='#0b1217',
                      fg='white', font=('', 10))
        hh.place(x=30, y=70)
        hh.bind('<Button-1>', open_modem)
        hs = tk.Label(ssss, text='Дополнительные', background='#0b1217',
                      fg='white', font=('', 10))
        hs.place(x=30, y=90)
        hs.bind('<Button-1>', open_dop)

    def close_pod_menu_detot(_):
        global atsosite
        lkj.place(x=10, y=300)
        sdf.place(x=10, y=360)
        hhhhg.place(x=65, y=310)
        hhhhh.place(x=65, y=370)
        ff2.place(x=250, y=310)
        ff3.place(x=250, y=370)
        atsosite.destroy()
        ff1.config(image=galka)

    def close_pod_menu_ychet(_):
        global ststs
        sdf.place(x=10, y=360)
        hhhhh.place(x=65, y=370)
        ff3.place(x=250, y=370)
        ststs.destroy()
        ff2.config(image=galka)

    def close_pod_menu_adnim(_):
        global ssss
        ssss.destroy()
        ff3.config(image=galka)

    open_main_str()
    caphass.destroy()
    root.config(background='#ccc')
    lala = tk.Label(root, background='white', width=160, height=3)
    lala.place(x=0, y=0)

    klk = tk.Label(root, background='#1e2329', width=40, height=50)
    klk.place(x=0, y=50)

    lkl = tk.Label(root, background='#060a0d', width=150, height=3)
    lkl.place(x=285, y=50)

    ddd = tk.Label(root, text='ООО Торговые Автоматы',
                   background='#060a0d', fg='white',
                   font=('', 15))
    ddd.place(x=290, y=60)

    sls = tk.Label(root, background='white', border=2, relief='solid',
                   width=20, height=3)
    sls.place(x=980, y=0)
    ttt = tk.Label(sls, text='Контаурова Е.С', background='white',
                   fg='black')
    ttt.place(x=40, y=5)
    tt = tk.Label(sls, text='Администратор', background='white',
                  fg='#7f8182')
    tt.place(x=40, y=20)
    one = tk.Label(sls, width=2, background='#c7c7c7')
    one.place(x=7, y=13)
    one = tk.Label(sls, width=2, background='#2809ed')
    one.place(x=7, y=18)
    one = tk.Label(sls, width=2, background='#ed1826')
    one.place(x=7, y=25)
    one = tk.Label(sls, width=2, background='#fff')
    one.place(x=7, y=31)
    one = tk.Label(sls, width=2, background='#060a0d')
    one.place(x=7, y=47)

    ksks = tk.Label(root, background='#060a0d', width=10, height=3)
    ksks.place(x=206, y=50)
    menu = PhotoImage(file='image/menu.png')
    menu = menu.subsample(5)
    dd = tk.Label(root, image=menu, background='#060a0d')
    dd.place(x=215, y=55)
    # asd = tk.Label(text='Личный кабинет. Главная', font=('', 15))
    # asd.place(x=300, y=110)

    asd = tk.Label(root, text='Навигация', background='#1e2329', fg='#c4cacf',
                   font=('', 13))
    asd.place(x=10, y=62)

    df = PhotoImage(file='image/frame12.png')
    df = df.subsample(6)
    ass = tk.Label(root, image=df, background='#1e2329')
    ass.place(x=10, y=120)
    hghg = tk.Label(root, text='Главная', font=('', 14), background='#1e2329',
                    fg='#c4cacf')
    hghg.place(x=65, y=130)
    hghg.bind('<Button-1>', lambda e: open_main_str())

    dfd = PhotoImage(file='image/frame13.png')
    dfd = dfd.subsample(6)
    asff = tk.Label(root, image=dfd, background='#1e2329')
    asff.place(x=10, y=180)
    hghgg = tk.Label(root, text='Монитор ТА', font=('', 14),
                     background='#1e2329',
                     fg='#c4cacf')
    hghgg.place(x=65, y=190)
    hghgg.bind('<Button-1>', open_monik_str)

    aaa = PhotoImage(file='image/frame14.png')
    aaa = aaa.subsample(6)
    asas = tk.Label(root, image=aaa, background='#1e2329')
    asas.place(x=10, y=240)
    hghhg = tk.Label(root, text='Детальные отчеты', font=('', 14),
                     background='#1e2329',
                     fg='#c4cacf')
    hghhg.place(x=65, y=250)
    hghhg.bind('<Button-1>', open_pod_menu_detot)

    ggg = PhotoImage(file='image/frame16.png')
    ggg = ggg.subsample(6)

    lkj = tk.Label(root, image=ggg, background='#1e2329')
    lkj.place(x=10, y=300)
    hhhhg = tk.Label(root, text='Учет ТМЦ', font=('', 14),
                     background='#1e2329',
                     fg='#c4cacf')
    hhhhg.place(x=65, y=310)
    hhhhg.bind('<Button-1>', open_pod_menu_ychet)

    ddd = PhotoImage(file='image/frame15.png')
    ddd = ddd.subsample(6)
    sdf = tk.Label(root, image=ddd, background='#1e2329')
    sdf.place(x=10, y=360)
    hhhhh = tk.Label(root, text='Администрирование', font=('', 14),
                     background='#1e2329',
                     fg='#c4cacf')
    hhhhh.place(x=65, y=370)
    hhhhh.bind('<Button-1>', open_pod_menu_admin)

    galka = PhotoImage(file='image/galka.png')
    galka = galka.subsample(10)

    soss = PhotoImage(file='image/menuk.png')
    soss = soss.subsample(10)

    ff1 = tk.Label(root, image=galka, background='#1e2329')
    ff1.place(x=250, y=250)
    ff1.bind('<Button-1>', close_pod_menu_detot)
    ff2 = tk.Label(root, image=galka, background='#1e2329')
    ff2.place(x=250, y=310)
    ff2.bind('<Button-1>', close_pod_menu_ychet)
    ff3 = tk.Label(root, image=galka, background='#1e2329')
    ff3.place(x=250, y=370)
    ff3.bind('<Button-1>', close_pod_menu_adnim)


def get_autoriz():
    aut = AutorizRegus.select()
    return aut


def avtoriza(event=None):
    global lala
    global caphass
    lala = tk.Label(root, background='#1c1c15', width=50,
                    height=30)
    lala.place(x=400, y=100)
    zarag = tk.Label(lala, text='АВТОРИЗАЦИЯ', fg='white',
                     background='#1c1c15', font=('', 20, 'bold'))
    zarag.place(x=70, y=15)
    user = tk.Label(lala, text='Логин: ', fg='white',
                    background='#1c1c15', font=('', 14))
    user.place(x=10, y=90)
    us = tk.Entry(lala)
    us.place(x=120, y=95)

    pasw = tk.Label(lala, text='Пароль: ', fg='white',
                    background='#1c1c15', font=('', 14))
    pasw.place(x=10, y=150)
    pa = tk.Entry(lala)
    pa.place(x=120, y=155)

    def proverka():
        pole_one = us.get()
        pole_two = pa.get()
        aut = get_autoriz()
        for i in aut:
            if pole_one == i.user and pole_two == i.password:
                # open_glavna()
                capha()
                # messagebox.showinfo('Успех', f'Добро пожаловать {i.user}!')
                break
        else:
            messagebox.showinfo('Ошибка', 'Такого пользователя не найдено')

    b = tk.Button(lala, text='Войти', command=proverka,
                  width=30, background='#6b6b5e', fg='white')
    b.place(x=65, y=380)

    regis = tk.Label(lala, text='Зарегестрироваться', background='#1c1c15',
                     fg='white')
    regis.place(x=120, y=423)
    regis.bind('<Button-1>', requsts)

    def capha():
        global caphass

        def upup():
            uu = main.get()
            if str(three) == str(uu):
                open_glavna()
            else:
                messagebox.showerror('Неверно', 'Попробуйте еще раз')

        caphass = tk.Toplevel()
        caphass.geometry('200x100')
        caphass.title('Капча')
        one = random.randint(1, 100)
        two = random.randint(1, 100)
        cc = random.randint(1, 100)
        three = one + two - cc
        print(three)
        cap1 = tk.Label(caphass, text=f'{one}+{two}-{cc}')
        cap1.pack()
        main = tk.Entry(caphass)
        main.pack()
        knopka = tk.Button(caphass, text='Отправить', command=upup)
        knopka.pack()


def requsts(_):
    global lalal
    lalal = tk.Label(root, background='#1c1c15', width=50,
                     height=30)
    lalal.place(x=400, y=100)
    zarag = tk.Label(lalal, text='РЕГИСТРАЦИЯ', fg='white',
                     background='#1c1c15', font=('', 20, 'bold'))
    zarag.place(x=70, y=15)
    user = tk.Label(lalal, text='Придумайте логин: ', fg='white',
                    background='#1c1c15', font=('', 14))
    user.place(x=10, y=90)
    us = tk.Entry(lalal)
    us.place(x=200, y=95)

    pasw = tk.Label(lalal, text='Придуймате пароль: ', fg='white',
                    background='#1c1c15', font=('', 14))
    pasw.place(x=10, y=150)
    pa = tk.Entry(lalal)
    pa.place(x=200, y=155)

    def reg():
        userget = us.get()
        paswget = pa.get()

        _ = AutorizRegus.create(
            user=userget,
            password=paswget
        )
        messagebox.showinfo('Успех', 'Вы успешно зарегестрировались!')
        capha()

    b = tk.Button(lalal, text='Зарегистрироваться', command=reg,
                  width=30, background='#6b6b5e', fg='white')
    b.place(x=65, y=380)

    avt = tk.Label(lalal, text='Авторизоваться', background='#1c1c15',
                   fg='white')
    avt.place(x=130, y=423)
    avt.bind('<Button-1>', lambda e: avtoriza())

    def capha():
        global caphass

        def upup():
            uu = main.get()
            if str(three) == str(uu):
                open_glavna()
            else:
                messagebox.showerror('Неверно', 'Попробуйте еще раз')

        caphass = tk.Toplevel()
        caphass.geometry('200x100')
        caphass.title('Капча')
        one = random.randint(1, 100)
        two = random.randint(1, 100)
        cc = random.randint(1, 100)
        three = one + two - cc
        print(three)
        cap1 = tk.Label(caphass, text=f'{one}+{two}-{cc}')
        cap1.pack()
        main = tk.Entry(caphass)
        main.pack()
        knopka = tk.Button(caphass, text='Отправить', command=upup)
        knopka.pack()


avtoriza()
root.mainloop()
